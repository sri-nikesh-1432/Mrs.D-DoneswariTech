"""
Document Processor — Extract and clean text from uploaded files.
Supports: PDF, DOCX, TXT, CSV
"""

import os
import re
import traceback
from typing import Optional
import pandas as pd

from app.utils.logger import get_logger

logger = get_logger(__name__)

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".csv"}
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB


def validate_file(filename: str, file_size: int) -> tuple[bool, str]:
    """Validate file type and size. Returns (is_valid, error_message)."""
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        return False, f"Unsupported file type: {ext}. Allowed: {', '.join(ALLOWED_EXTENSIONS)}"
    if file_size > MAX_FILE_SIZE:
        return False, f"File too large: {file_size / 1024 / 1024:.1f} MB. Max: 50 MB"
    if file_size == 0:
        return False, "File is empty"
    return True, ""


async def extract_text(file_path: str, filename: str) -> str:
    """
    Extract text from an uploaded file based on its extension.
    Returns the extracted plain text.
    """
    ext = os.path.splitext(filename)[1].lower()
    logger.info("Extracting text from %s (%s)", filename, ext)

    try:
        if ext == ".pdf":
            text = _extract_pdf(file_path)
        elif ext == ".docx":
            text = _extract_docx(file_path)
        elif ext == ".txt":
            text = _extract_txt(file_path)
        elif ext == ".csv":
            text = _extract_csv(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

        cleaned = clean_text(text)
        logger.info("Extracted %d characters from %s", len(cleaned), filename)
        return cleaned

    except Exception as e:
        logger.error("Failed to extract text from %s: %s\n%s", filename, e, traceback.format_exc())
        raise


def _extract_pdf(file_path: str) -> str:
    """Extract text from PDF using PyMuPDF."""
    import fitz
    doc = fitz.open(file_path)
    text_parts = []
    for page_num, page in enumerate(doc):
        page_text = page.get_text()
        if page_text.strip():
            text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
    doc.close()
    return "\n\n".join(text_parts)


def _extract_docx(file_path: str) -> str:
    """Extract text from DOCX using python-docx."""
    from docx import Document
    doc = Document(file_path)
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            if row_text.strip():
                text_parts.append(row_text)
    return "\n".join(text_parts)


def _extract_txt(file_path: str) -> str:
    """Extract text from TXT file."""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _extract_csv(file_path: str) -> str:
    """Extract text from CSV file using pandas."""
    df = pd.read_csv(file_path)
    text_parts = []
    # Add column headers
    text_parts.append(" | ".join(str(col) for col in df.columns))
    # Add rows
    for _, row in df.iterrows():
        text_parts.append(" | ".join(str(val) for val in row))
    return "\n".join(text_parts)


def clean_text(text: str) -> str:
    """
    Clean extracted text by removing artifacts.
    """
    # Remove null bytes
    text = text.replace("\x00", "")
    # Normalize line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Remove excessive blank lines (keep max 2)
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Remove headers/footers (page numbers)
    text = re.sub(r"\n\s*\d+\s*\n", "\n", text)
    # Remove URLs
    text = re.sub(r"https?://\S+", "", text)
    # Normalize Unicode
    import unicodedata
    text = unicodedata.normalize("NFKC", text)
    # Strip leading/trailing whitespace per line
    lines = [line.strip() for line in text.split("\n")]
    text = "\n".join(lines)
    # Remove empty lines at start/end
    text = text.strip()
    return text
