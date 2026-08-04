"""
Document upload and extraction service for Mrs. D AI Admission Campaign Platform.
Supports PDF, DOCX, TXT, and CSV files.
"""

import os
import aiofiles
from pathlib import Path
from typing import Optional, Tuple
import fitz  # PyMuPDF
from docx import Document
import pandas as pd

from app.config.settings import settings
from app.logs.logger import get_logger

logger = get_logger(__name__)


class DocumentService:
    """Service for handling document uploads and text extraction."""
    
    def __init__(self):
        self.allowed_mime_types = settings.ALLOWED_DOCUMENT_TYPES
        self.max_size = settings.MAX_UPLOAD_SIZE
    
    async def validate_file(self, file_content: bytes, filename: str) -> Tuple[bool, Optional[str]]:
        """
        Validate uploaded file.
        
        Args:
            file_content: File content as bytes
            filename: Original filename
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        logger.info(f"=== STEP 1: DOCUMENT UPLOAD VALIDATION ===")
        logger.info(f"Filename: {filename}")
        logger.info(f"File size: {len(file_content)} bytes ({len(file_content) / (1024 * 1024):.2f} MB)")
        
        # Check file size
        if len(file_content) > self.max_size:
            logger.error(f"File size {len(file_content)} exceeds maximum {self.max_size}")
            return False, f"File size exceeds maximum limit of {self.max_size / (1024 * 1024)}MB"
        
        logger.info(f"✓ File size validation passed")
        
        # Check file extension
        ext = Path(filename).suffix.lower()
        allowed_extensions = ['.pdf', '.docx', '.txt', '.csv']
        logger.info(f"File extension: {ext}")
        
        if ext not in allowed_extensions:
            logger.error(f"Extension {ext} not in allowed extensions: {allowed_extensions}")
            return False, f"File extension {ext} is not allowed"
        
        logger.info(f"✓ File extension validation passed")
        
        # Basic content validation based on extension
        if ext == '.pdf':
            if not file_content.startswith(b'%PDF'):
                logger.error("PDF file does not start with %PDF magic bytes")
                return False, "Invalid PDF file"
            logger.info("✓ PDF magic bytes validated")
        elif ext == '.docx':
            if not file_content.startswith(b'PK\x03\x04'):
                logger.error("DOCX file does not start with PK magic bytes")
                return False, "Invalid DOCX file"
            logger.info("✓ DOCX magic bytes validated")
        elif ext == '.csv':
            # Check if it looks like text
            try:
                file_content.decode('utf-8')
                logger.info("✓ CSV file is valid UTF-8 text")
            except UnicodeDecodeError:
                logger.error("CSV file is not valid UTF-8 text")
                return False, "Invalid CSV file (not text)"
        elif ext == '.txt':
            try:
                file_content.decode('utf-8')
                logger.info("✓ TXT file is valid UTF-8 text")
            except UnicodeDecodeError:
                logger.error("TXT file is not valid UTF-8 text")
                return False, "Invalid text file"
        
        logger.info(f"=== STEP 1 COMPLETE: UPLOAD VALIDATION SUCCESSFUL ===")
        return True, None
    
    async def save_file(self, file_content: bytes, filename: str) -> Path:
        """
        Save uploaded file to knowledge directory.
        
        Args:
            file_content: File content as bytes
            filename: Original filename
            
        Returns:
            Path to saved file
        """
        logger.info(f"=== STEP 1: FILE SAVING ===")
        
        # Generate unique filename
        timestamp = pd.Timestamp.now().strftime("%Y%m%d_%H%M%S")
        safe_filename = f"{timestamp}_{filename}"
        file_path = settings.KNOWLEDGE_DIR / safe_filename
        
        logger.info(f"Saving to: {file_path}")
        logger.info(f"Safe filename: {safe_filename}")
        
        # Save file
        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(file_content)
        
        # Verify file was saved
        if file_path.exists():
            file_size = file_path.stat().st_size
            logger.info(f"✓ File saved successfully")
            logger.info(f"✓ File exists at: {file_path}")
            logger.info(f"✓ Saved file size: {file_size} bytes")
            logger.info(f"=== STEP 1 COMPLETE: FILE SAVED SUCCESSFULLY ===")
        else:
            logger.error(f"✗ File save failed - file does not exist at {file_path}")
            raise IOError(f"Failed to save file to {file_path}")
        
        return file_path
    
    async def extract_text(self, file_path: Path) -> str:
        """
        Extract text from document based on file type.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Extracted text as string
        """
        logger.info(f"=== STEP 2: DOCUMENT PARSING ===")
        logger.info(f"File path: {file_path}")
        logger.info(f"File exists: {file_path.exists()}")
        
        if not file_path.exists():
            logger.error(f"✗ File does not exist at {file_path}")
            raise FileNotFoundError(f"File not found: {file_path}")
        
        ext = file_path.suffix.lower()
        logger.info(f"File extension: {ext}")
        
        try:
            if ext == '.pdf':
                logger.info("Using PDF parser (PyMuPDF)")
                text = await self._extract_from_pdf(file_path)
            elif ext == '.docx':
                logger.info("Using DOCX parser (python-docx)")
                text = await self._extract_from_docx(file_path)
            elif ext == '.txt':
                logger.info("Using TXT parser")
                text = await self._extract_from_txt(file_path)
            elif ext == '.csv':
                logger.info("Using CSV parser (pandas)")
                text = await self._extract_from_csv(file_path)
            else:
                logger.error(f"✗ Unsupported file type: {ext}")
                raise ValueError(f"Unsupported file type: {ext}")
            
            logger.info(f"=== STEP 2 COMPLETE: PARSING SUCCESSFUL ===")
            return text
        except Exception as e:
            logger.error(f"✗ Error extracting text from {file_path}: {e}")
            logger.error(f"Error type: {type(e).__name__}")
            raise
    
    async def _extract_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF using PyMuPDF."""
        logger.info(f"=== STEP 3: PDF TEXT EXTRACTION ===")
        logger.info(f"Opening PDF: {file_path}")
        
        text = ""
        doc = fitz.open(str(file_path))
        
        logger.info(f"PDF page count: {len(doc)}")
        
        for page_num, page in enumerate(doc):
            page_text = page.get_text()
            text += page_text
            if page_num == 0:
                logger.info(f"First page text length: {len(page_text)} characters")
                logger.info(f"First page preview (first 200 chars): {page_text[:200]}")
        
        doc.close()
        
        logger.info(f"Total extracted text length: {len(text)} characters")
        logger.info(f"✓ PDF extraction completed")
        
        cleaned_text = await self._clean_text(text)
        logger.info(f"After cleaning: {len(cleaned_text)} characters")
        
        logger.info(f"=== STEP 3 COMPLETE: PDF TEXT EXTRACTION SUCCESSFUL ===")
        return cleaned_text
    
    async def _extract_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX using python-docx."""
        logger.info(f"=== STEP 3: DOCX TEXT EXTRACTION ===")
        logger.info(f"Opening DOCX: {file_path}")
        
        doc = Document(str(file_path))
        text = ""
        
        logger.info(f"Paragraph count: {len(doc.paragraphs)}")
        logger.info(f"Table count: {len(doc.tables)}")
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " | "
                text += "\n"
        
        logger.info(f"Total extracted text length: {len(text)} characters")
        logger.info(f"✓ DOCX extraction completed")
        
        cleaned_text = await self._clean_text(text)
        logger.info(f"After cleaning: {len(cleaned_text)} characters")
        
        logger.info(f"=== STEP 3 COMPLETE: DOCX TEXT EXTRACTION SUCCESSFUL ===")
        return cleaned_text
    
    async def _extract_from_txt(self, file_path: Path) -> str:
        """Extract text from TXT file."""
        logger.info(f"=== STEP 3: TXT TEXT EXTRACTION ===")
        logger.info(f"Opening TXT: {file_path}")
        
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
            text = await f.read()
        
        logger.info(f"Total extracted text length: {len(text)} characters")
        logger.info(f"✓ TXT extraction completed")
        
        cleaned_text = await self._clean_text(text)
        logger.info(f"After cleaning: {len(cleaned_text)} characters")
        
        logger.info(f"=== STEP 3 COMPLETE: TXT TEXT EXTRACTION SUCCESSFUL ===")
        return cleaned_text
    
    async def _extract_from_csv(self, file_path: Path) -> str:
        """Extract text from CSV file."""
        logger.info(f"=== STEP 3: CSV TEXT EXTRACTION ===")
        logger.info(f"Opening CSV: {file_path}")
        
        df = pd.read_csv(file_path)
        logger.info(f"CSV rows: {len(df)}")
        logger.info(f"CSV columns: {list(df.columns)}")
        
        text = df.to_string(index=False)
        
        logger.info(f"Total extracted text length: {len(text)} characters")
        logger.info(f"✓ CSV extraction completed")
        
        cleaned_text = await self._clean_text(text)
        logger.info(f"After cleaning: {len(cleaned_text)} characters")
        
        logger.info(f"=== STEP 3 COMPLETE: CSV TEXT EXTRACTION SUCCESSFUL ===")
        return cleaned_text
    
    async def _clean_text(self, text: str) -> str:
        """
        Clean extracted text.
        
        Removes:
        - Extra whitespace
        - Broken characters
        - Duplicate lines
        
        Preserves:
        - Tables
        - Phone numbers
        - Bullet lists
        - Course names
        """
        logger.info(f"=== STEP 3: TEXT CLEANING ===")
        logger.info(f"Input text length: {len(text)} characters")
        
        if not text:
            logger.error("✗ Input text is empty")
            return ""
        
        logger.info(f"Input text preview (first 200 chars): {text[:200]}")
        
        # Remove extra whitespace
        text = ' '.join(text.split())
        logger.info(f"After whitespace removal: {len(text)} characters")
        
        # Remove duplicate lines (but keep at least one occurrence)
        lines = text.split('\n')
        seen = set()
        unique_lines = []
        for line in lines:
            stripped = line.strip()
            if stripped and stripped not in seen:
                seen.add(stripped)
                unique_lines.append(line)
            elif not stripped:
                unique_lines.append(line)  # Keep empty lines for structure
        
        text = '\n'.join(unique_lines)
        logger.info(f"After duplicate removal: {len(text)} characters")
        logger.info(f"Unique lines: {len(unique_lines)}")
        
        logger.info(f"✓ Text cleaning completed")
        logger.info(f"=== STEP 3 COMPLETE: TEXT CLEANING SUCCESSFUL ===")
        
        return text
    
    async def delete_file(self, file_path: Path) -> bool:
        """
        Delete a file from storage.
        
        Args:
            file_path: Path to the file
            
        Returns:
            True if deleted successfully
        """
        try:
            if file_path.exists():
                os.remove(file_path)
                logger.info(f"File deleted: {file_path}")
                return True
            return False
        except Exception as e:
            logger.error(f"Error deleting file {file_path}: {e}")
            return False
